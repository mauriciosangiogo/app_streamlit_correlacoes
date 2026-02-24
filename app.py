import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from scipy import stats
from scipy.stats import shapiro
import seaborn as sns
from io import BytesIO
import matplotlib
import docx

matplotlib.rcParams['font.family'] = 'DejaVu Sans'

# ── Configuração da página ──────────────────────────────────────────────────
st.set_page_config(page_title="Análise Estatística", page_icon="📊", layout="wide")

# ── CSS ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Source+Sans+3:wght@300;400;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Source Sans 3', sans-serif; }
    .block-container { padding-top: 2.2rem; }
    .main-title { font-size: 1.8rem; font-weight: 700; color: #1a1a2e; margin-bottom: 0.1rem; }
    .subtitle { font-size: 1rem; color: #6c757d; margin-bottom: 1.2rem; }
    .section-header { font-size: 1.05rem; font-weight: 600; color: #1a1a2e; border-bottom: 2px solid #e63946; padding-bottom: 0.3rem; margin-bottom: 1rem; }
    .metric-card { background: #f8f9fa; border-radius: 8px; padding: 0.8rem 1rem; border-left: 4px solid #e63946; margin-bottom: 0.5rem; }
    .metric-label { font-size: 0.75rem; color: #6c757d; text-transform: uppercase; letter-spacing: 0.5px; }
    .metric-value { font-size: 1.3rem; font-weight: 700; color: #1a1a2e; }
    div[data-testid="stSidebar"] { background: #f8f9fa; }
    .info-box { background: #f0f4f8; border-radius: 8px; padding: 1rem 1.2rem; margin: 0.5rem 0; border-left: 3px solid #457b9d; font-size: 0.9rem; color: #1d3557; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# AUTENTICAÇÃO
# ══════════════════════════════════════════════════════════════════════════════
def verificar_credenciais(usuario, senha):
    try:
        senhas = st.secrets["passwords"]
        if usuario in senhas and senhas[usuario] == senha:
            return True
    except (KeyError, FileNotFoundError):
        st.error("⚠️ Secrets não configurados. Configure [passwords] nos Secrets do Streamlit Cloud.")
    return False

def tela_login():
    st.markdown("<style>section[data-testid='stSidebar']{display:none;}</style>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 1.2, 1])
    with col2:
        st.markdown("")
        st.markdown('<div class="main-title" style="text-align:center;">📊 Análise Estatística</div>', unsafe_allow_html=True)
        st.markdown('<div class="subtitle" style="text-align:center;">Faça login para acessar</div>', unsafe_allow_html=True)
        with st.form("login_form"):
            usuario = st.text_input("Usuário")
            senha = st.text_input("Senha", type="password")
            submit = st.form_submit_button("Entrar", use_container_width=True)
            if submit:
                if verificar_credenciais(usuario, senha):
                    st.session_state["autenticado"] = True
                    st.session_state["usuario"] = usuario
                    st.rerun()
                else:
                    st.error("Usuário ou senha incorretos.")

if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = False
if not st.session_state["autenticado"]:
    tela_login()
    st.stop()

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ══════════════════════════════════════════════════════════════════════════════
CORES_DISPONIVEIS = {
    'Azul': 'tab:blue', 'Laranja': 'tab:orange', 'Verde': 'tab:green',
    'Vermelho': 'tab:red', 'Roxo': 'tab:purple', 'Marrom': 'tab:brown',
    'Rosa': 'tab:pink', 'Cinza': 'tab:gray', 'Oliva': 'tab:olive',
    'Ciano': 'tab:cyan', 'Azul claro': 'lightblue', 'Pêssego': 'peachpuff',
    'Verde claro': 'lightgreen', 'Salmão': 'salmon', 'Ameixa': 'plum',
    'Trigo': 'wheat', 'Rosa claro': 'lightpink', 'Prata': 'silver',
    'Amarelo-verde': 'yellowgreen', 'Turquesa': 'paleturquoise'
}

# ══════════════════════════════════════════════════════════════════════════════
# FUNÇÕES
# ══════════════════════════════════════════════════════════════════════════════

def tabela_cientifica(df_tabela, titulo=""):
    """Tabela em formato científico: bordas top, abaixo do header e bottom apenas."""
    fig_t, ax = plt.subplots(figsize=(max(8, len(df_tabela.columns) * 1.8),
                                       max(2, len(df_tabela) * 0.45 + 1.2)))
    ax.axis('off')

    cell_text = []
    for _, row in df_tabela.iterrows():
        formatted = []
        for val in row:
            if isinstance(val, (int, np.integer)):
                formatted.append(f"{val:,}")
            elif isinstance(val, (float, np.floating)):
                formatted.append(f"{val:.4f}" if pd.notna(val) else "—")
            else:
                formatted.append(str(val))
        cell_text.append(formatted)

    table = ax.table(
        cellText=cell_text,
        colLabels=df_tabela.columns.tolist(),
        rowLabels=df_tabela.index.tolist(),
        cellLoc='center', rowLoc='center', loc='center'
    )
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 1.6)

    n_rows = len(df_tabela)
    for key, cell in table.get_celld().items():
        row, col = key
        cell.set_facecolor('white')
        cell.set_edgecolor('white')
        cell.set_text_props(fontfamily='DejaVu Sans')
        cell.visible_edges = ''

        if row == 0:  # header
            cell.set_text_props(fontweight='bold', fontsize=9)
            cell.set_edgecolor('black')
            cell.visible_edges = 'BT'
            cell.set_linewidth(1.2)
        elif row == n_rows:  # última linha
            cell.set_edgecolor('black')
            cell.visible_edges = 'B'
            cell.set_linewidth(1.2)

        if col == -1:  # row labels
            cell.set_text_props(fontweight='bold', fontsize=9, ha='right')
            if row == 0:
                cell.visible_edges = 'BT'
                cell.set_edgecolor('black')
                cell.set_linewidth(1.2)
            elif row == n_rows:
                cell.visible_edges = 'B'
                cell.set_edgecolor('black')
                cell.set_linewidth(1.2)
            else:
                cell.visible_edges = ''

    if titulo:
        ax.set_title(titulo, fontsize=11, fontweight='bold', pad=15, loc='left')
    fig_t.patch.set_facecolor('white')
    fig_t.tight_layout()
    return fig_t


def criar_grafico_barras(df_plot, var_cat, var_num, titulo, label_x, label_y, cor, agregacao='sum'):
    # Validação: variáveis não podem ser iguais
    if var_cat == var_num:
        raise ValueError("A variável categórica e a variável numérica devem ser diferentes.")
    
    df_ag = df_plot.groupby(var_cat, as_index=False)[var_num].agg(agregacao)
    df_ag = df_ag.sort_values(by=var_cat, ascending=True)
    vx = df_ag[var_cat].astype(str)
    vy = df_ag[var_num]

    fig, ax = plt.subplots(figsize=(10, max(4, len(vx) * 0.5 + 1)))
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    ax.grid(False)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(left=False, bottom=False, right=False, top=False)

    bars = ax.barh(vx, vy, color=cor, edgecolor='black', linewidth=1.5, alpha=0.8)
    for bar in bars:
        w = bar.get_width()
        label = f'{int(w):,}' if w == int(w) else f'{w:,.2f}'
        ax.text(w + max(vy) * 0.02, bar.get_y() + bar.get_height() / 2,
                label, ha='left', va='center', fontsize=11, fontweight='bold', color='black')

    ax.set_xlabel(label_x, fontsize=12, fontweight='bold', color='black')
    ax.set_ylabel(label_y, fontsize=12, fontweight='bold', color='black')
    ax.set_title(titulo, fontsize=14, fontweight='bold', pad=20, color='black')
    ax.tick_params(axis='x', colors='black')
    ax.tick_params(axis='y', colors='black')
    ax.set_xlim(0, max(vy) * 1.15 if len(vy) > 0 else 1)
    fig.tight_layout()
    return fig


def calcular_correlacao(df, vars_x, vars_y, metodo):
    func = {"pearson": stats.pearsonr, "spearman": stats.spearmanr, "kendall": stats.kendalltau}[metodo]
    mat_r = pd.DataFrame(index=vars_x, columns=vars_y, dtype=float)
    mat_p = pd.DataFrame(index=vars_x, columns=vars_y, dtype=float)
    mat_n = pd.DataFrame(index=vars_x, columns=vars_y, dtype=int)
    for x in vars_x:
        for y in vars_y:
            v = df[[x, y]].dropna()
            n = len(v)
            mat_n.loc[x, y] = n
            if n > 2:
                r, p = func(v[x], v[y])
                mat_r.loc[x, y] = r
                mat_p.loc[x, y] = p
            else:
                mat_r.loc[x, y] = np.nan
                mat_p.loc[x, y] = np.nan
    return mat_r, mat_p, mat_n


def formatar_tabela_corr(mat_r, mat_p, alpha):
    tab = pd.DataFrame(index=mat_r.index, columns=mat_r.columns)
    for x in mat_r.index:
        for y in mat_r.columns:
            r, p = mat_r.loc[x, y], mat_p.loc[x, y]
            if pd.notna(p):
                if p < 0.001: s = "***"
                elif p < 0.01: s = "**"
                elif p < alpha: s = "*"
                else: s = "ns"
                tab.loc[x, y] = f"{r:.3f} {s}"
            else:
                tab.loc[x, y] = "NA"
    return tab


def criar_heatmap(mat_r, mat_p, alpha, paleta, metodo, tit_x, tit_y):
    fig, ax = plt.subplots(figsize=(max(6, len(mat_r.columns) * 1.6), max(4, len(mat_r.index) * 1.1)))
    annot = mat_r.copy().astype(object)
    for x in mat_r.index:
        for y in mat_r.columns:
            r, p = mat_r.loc[x, y], mat_p.loc[x, y]
            if pd.notna(r) and pd.notna(p):
                if p < 0.001: s = "***"
                elif p < 0.01: s = "**"
                elif p < alpha: s = "*"
                else: s = ""
                annot.loc[x, y] = f"{r:.2f}{s}"
            else:
                annot.loc[x, y] = "NA"

    nomes = {"pearson": "Pearson (r)", "spearman": "Spearman (ρ)", "kendall": "Kendall (τ)"}
    sns.heatmap(mat_r.astype(float), annot=annot, fmt="", cmap=paleta,
            center=0, vmin=-1, vmax=1, square=True,
            linewidths=0.8, linecolor="#ffffff",
            cbar_kws={"shrink": 0.8},
            ax=ax, annot_kws={"size": 6})

# Ajustar fontes da colorbar
    cbar = ax.collections[0].colorbar
    cbar.ax.tick_params(labelsize=6)
    cbar.set_label(f"Correlação ({metodo.capitalize()})", fontsize=6)
    ax.set_title(f"Matriz de Correlação — {nomes[metodo]}", fontsize=9, fontweight="bold", pad=14)
    ax.set_xlabel(tit_x, fontsize=6, fontweight="bold")
    ax.set_ylabel(tit_y, fontsize=6, fontweight="bold")
    ax.tick_params(axis='x', rotation=45, labelsize=6)
    ax.tick_params(axis='y', rotation=0, labelsize=6)
    fig.text(0.5, -0.02, f"*** p<0.001   ** p<0.01   * p<{alpha}   (sem asterisco = não significativo)",
             ha="center", fontsize=6, color="#6c757d")
    fig.patch.set_facecolor('white')
    fig.tight_layout()
    return fig


def fig_to_bytes(fig, dpi=300):
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight", facecolor="white")
    buf.seek(0)
    return buf


def gerar_word(figs_dict, metodo, alpha):
    """Gera relatório Word com python-docx."""
    from docx import Document as DocxDoc
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    p = doc.add_heading('Relatório de Análise Estatística', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    nomes_m = {"pearson": "Pearson", "spearman": "Spearman", "kendall": "Kendall"}

    secoes = [
        ("1. Estatísticas Descritivas", "desc",
         "Tabela com medidas de tendência central e dispersão das variáveis numéricas."),
        ("2. Teste de Normalidade (Shapiro-Wilk)", "shapiro",
         "Avalia se os dados seguem distribuição normal. p > 0.05 indica normalidade."),
        ("3. Gráfico Descritivo", "barras", ""),
        (f"4. Heatmap de Correlação ({nomes_m.get(metodo, metodo)})", "heatmap",
         f"Método: {nomes_m.get(metodo, metodo)}. α = {alpha}. *** p<0.001, ** p<0.01, * p<{alpha}."),
        ("5. Tabela de Correlações", "corr_tab", ""),
    ]

    for titulo, chave, descricao in secoes:
        if chave in figs_dict and figs_dict[chave] is not None:
            doc.add_heading(titulo, level=2)
            if descricao:
                doc.add_paragraph(descricao)
            buf = fig_to_bytes(figs_dict[chave])
            doc.add_picture(buf, width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Interpretação
    doc.add_heading('6. Interpretação da Força da Correlação', level=2)
    for linha in ["|r| < 0.10: Muito fraca", "0.10 ≤ |r| < 0.30: Fraca",
                  "0.30 ≤ |r| < 0.50: Moderada", "0.50 ≤ |r| < 0.70: Forte",
                  "|r| ≥ 0.70: Muito forte"]:
        doc.add_paragraph(linha)

    buf_doc = BytesIO()
    doc.save(buf_doc)
    buf_doc.seek(0)
    return buf_doc


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — DADOS
# ══════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="main-title">📊 Análise Estatística</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Descritivas · Gráficos · Correlações · Exportação</div>', unsafe_allow_html=True)

with st.sidebar:
    st.markdown(f"👤 **{st.session_state['usuario']}**")
    if st.button("Sair", use_container_width=True):
        st.session_state["autenticado"] = False
        st.session_state["usuario"] = ""
        st.rerun()
    st.markdown("---")
    st.markdown('<div class="section-header">📁 Dados</div>', unsafe_allow_html=True)
    arquivo = st.file_uploader("Carregar Excel", type=["xlsx", "xls"])
    if arquivo is None:
        st.info("Envie um arquivo Excel para começar.")
        st.stop()
    try:
        df = pd.read_excel(arquivo)
        st.success(f"{df.shape[0]} linhas × {df.shape[1]} colunas")
    except Exception as e:
        st.error(f"Erro: {e}")
        st.stop()
    cols_num = df.select_dtypes(include=[np.number]).columns.tolist()
    cols_todas = df.columns.tolist()
    if len(cols_num) < 2:
        st.error("Necessário ao menos 2 colunas numéricas.")
        st.stop()

# Armazenar figuras para export
if "figs" not in st.session_state:
    st.session_state["figs"] = {}

# ══════════════════════════════════════════════════════════════════════════════
# ABAS
# ══════════════════════════════════════════════════════════════════════════════
tab_desc, tab_graf, tab_corr, tab_export = st.tabs([
    "📋 Descritivas", "📊 Gráfico de Barras", "🔥 Correlações", "📥 Exportar"
])

# ── ABA 1: DESCRITIVAS ─────────────────────────────────────────────────────
with tab_desc:
    st.markdown('<div class="section-header">Estatísticas Descritivas</div>', unsafe_allow_html=True)
    cols_desc = st.multiselect("Selecione as variáveis", cols_num, default=cols_num[:5], key="dv")

    if cols_desc:
        desc_df = df[cols_desc].describe().round(4)
        fig_desc = tabela_cientifica(desc_df, titulo="Estatísticas Descritivas")
        st.pyplot(fig_desc, use_container_width=False)
        st.session_state["figs"]["desc"] = fig_desc

        buf_d = fig_to_bytes(fig_desc)
        st.download_button("⬇️ Baixar tabela descritivas (PNG)", buf_d,
                           file_name="descritivas.png", mime="image/png")

        with st.expander("ℹ️ Interpretação das estatísticas descritivas"):
            st.markdown("""
<div class="info-box">
<b>count:</b> Dados válidos (exclui NaN). <b>mean:</b> Média aritmética. <b>std:</b> Desvio padrão — dispersão em torno da média.<br>
<b>min/max:</b> Extremos. <b>25%, 50%, 75%:</b> Quartis (Q1, mediana, Q3).<br><br>
<b>Dicas:</b> média ≈ mediana → simétrica · média > mediana → assimetria positiva · média < mediana → assimetria negativa · IQR = Q3 − Q1.
</div>""", unsafe_allow_html=True)

        # Shapiro
        st.markdown("---")
        st.markdown('<div class="section-header">Teste de Normalidade (Shapiro-Wilk)</div>', unsafe_allow_html=True)
        alpha_sh = st.select_slider("α (Shapiro)", options=[0.01, 0.05, 0.10], value=0.05, key="ash")

        res_sh = {}
        for c in cols_desc:
            d = df[c].dropna()
            if len(d) >= 3:
                w, p = shapiro(d)
                res_sh[c] = {'W': round(w, 4), 'p-valor': round(p, 4), 'Normal?': 'Sim' if p > alpha_sh else 'Não'}
        if res_sh:
            df_sh = pd.DataFrame(res_sh).T
            fig_shapiro = tabela_cientifica(df_sh, titulo="Teste de Shapiro-Wilk")
            st.pyplot(fig_shapiro, use_container_width=False)
            st.session_state["figs"]["shapiro"] = fig_shapiro

            buf_sh = fig_to_bytes(fig_shapiro)
            st.download_button("⬇️ Baixar Shapiro (PNG)", buf_sh, file_name="shapiro.png", mime="image/png")

            with st.expander("ℹ️ Interpretação do teste de Shapiro-Wilk"):
                st.markdown(f"""
<div class="info-box">
<b>W:</b> Entre 0 e 1 — quanto mais perto de 1, mais normal. <b>p-valor:</b> Se p > {alpha_sh}, distribuição normal (não rejeita H₀).<br><br>
<b>Observações:</b> Confiável para n entre 3 e 5.000. Amostras muito grandes detectam desvios mínimos.
</div>""", unsafe_allow_html=True)
    else:
        st.warning("Selecione ao menos uma variável.")

# ── ABA 2: GRÁFICO DE BARRAS ───────────────────────────────────────────────
with tab_graf:
    st.markdown('<div class="section-header">Gráfico de Barras Horizontal</div>', unsafe_allow_html=True)

    c1, c2 = st.columns(2)
    with c1:
        var_cat = st.selectbox("Variável de agrupamento (eixo Y)", cols_todas, key="by")
        var_val = st.selectbox("Variável numérica (eixo X)", cols_num, key="bx")
        agg = st.selectbox("Agregação", ["sum", "mean"], format_func=lambda x: {"sum": "Soma", "mean": "Média"}[x], key="ba")
    with c2:
        tit_b = st.text_input("Título", value=f"Distribuição de {var_val} por {var_cat}", key="bt")
        lab_x = st.text_input("Rótulo eixo X", value=var_val, key="blx")
        lab_y = st.text_input("Rótulo eixo Y", value="", key="bly")
        cor_nome = st.selectbox("Cor das barras", list(CORES_DISPONIVEIS.keys()), index=0, key="bc")

    if var_cat == var_val:
        st.error("⚠️ A variável categórica e a variável numérica não podem ser a mesma coluna. Selecione colunas diferentes.")
    else:
        fig_barras = criar_grafico_barras(df, var_cat, var_val, tit_b, lab_x, lab_y, CORES_DISPONIVEIS[cor_nome], agg)
        st.pyplot(fig_barras, use_container_width=False)
        st.session_state["figs"]["barras"] = fig_barras

        buf_b = fig_to_bytes(fig_barras)
        st.download_button("⬇️ Baixar gráfico (PNG 300dpi)", buf_b, file_name="grafico_barras.png", mime="image/png")

# ── ABA 3: CORRELAÇÕES ─────────────────────────────────────────────────────
with tab_corr:
    st.markdown('<div class="section-header">Análise de Correlação</div>', unsafe_allow_html=True)

    cc1, cc2 = st.columns(2)
    with cc1:
        metodo = st.selectbox("Método", ["pearson", "spearman", "kendall"],
                              format_func=lambda x: {"pearson": "Pearson (paramétrico)", "spearman": "Spearman (não-paramétrico)", "kendall": "Kendall (amostras pequenas)"}[x], key="cm")
        vars_x = st.multiselect("Variáveis independentes (X)", cols_num, key="cx")
    with cc2:
        alpha_c = st.select_slider("α (correlação)", options=[0.01, 0.05, 0.10], value=0.05, key="ac")
        vars_y = st.multiselect("Variáveis dependentes (Y)", [c for c in cols_num if c not in vars_x], key="cy")

    if vars_x and vars_y:
        ct1, ct2 = st.columns(2)
        with ct1:
            hm_x = st.text_input("Título eixo X (heatmap)", "Variáveis Dependentes", key="hx")
        with ct2:
            hm_y = st.text_input("Título eixo Y (heatmap)", "Variáveis Independentes", key="hy")
        paleta = st.selectbox("Paleta", ["RdBu_r", "coolwarm", "PiYG", "BrBG", "PRGn", "RdYlGn"], key="cp")

        mat_r, mat_p, mat_n = calcular_correlacao(df, vars_x, vars_y, metodo)
        tab_fmt = formatar_tabela_corr(mat_r, mat_p, alpha_c)

        n_sig = ((mat_p.astype(float) < alpha_c) & mat_p.notna()).sum().sum()
        n_par = mat_r.notna().sum().sum()
        r_max = mat_r.astype(float).abs().max().max() if mat_r.notna().any().any() else 0

        m1, m2, m3 = st.columns(3)
        with m1:
            st.markdown(f'<div class="metric-card"><div class="metric-label">Pares</div><div class="metric-value">{n_par}</div></div>', unsafe_allow_html=True)
        with m2:
            st.markdown(f'<div class="metric-card"><div class="metric-label">Significativos</div><div class="metric-value">{n_sig}</div></div>', unsafe_allow_html=True)
        with m3:
            st.markdown(f'<div class="metric-card"><div class="metric-label">Maior |r|</div><div class="metric-value">{r_max:.3f}</div></div>', unsafe_allow_html=True)

        st.markdown("")
        fig_hm = criar_heatmap(mat_r, mat_p, alpha_c, paleta, metodo, hm_x, hm_y)
        st.pyplot(fig_hm, use_container_width=False)
        st.session_state["figs"]["heatmap"] = fig_hm
        buf_hm = fig_to_bytes(fig_hm)
        st.download_button("⬇️ Baixar heatmap (PNG 300dpi)", buf_hm, file_name="heatmap.png", mime="image/png")

        st.markdown("---")
        st.markdown("**Tabela de correlações**")
        fig_ct = tabela_cientifica(tab_fmt, titulo=f"Correlações ({metodo.capitalize()}) — *** p<0.001, ** p<0.01, * p<{alpha_c}")
        st.pyplot(fig_ct, use_container_width=False)
        st.session_state["figs"]["corr_tab"] = fig_ct
        buf_ct = fig_to_bytes(fig_ct)
        st.download_button("⬇️ Baixar tabela (PNG)", buf_ct, file_name="tabela_corr.png", mime="image/png")

        with st.expander("ℹ️ Força da correlação"):
            st.markdown("""
<div class="info-box">
|r| < 0.10: Muito fraca · 0.10–0.30: Fraca · 0.30–0.50: Moderada · 0.50–0.70: Forte · ≥ 0.70: Muito forte
</div>""", unsafe_allow_html=True)

        with st.expander("ℹ️ Escolha do método"):
            st.markdown("""
<div class="info-box">
<b>Pearson:</b> Relação linear, dados normais, contínuos, sem outliers extremos.<br>
<b>Spearman:</b> Quando Pearson não atende aos pressupostos. Usa ranks, avalia relação monotônica.<br>
<b>Kendall:</b> Amostras pequenas ou muitos empates nos dados.
</div>""", unsafe_allow_html=True)
    else:
        st.warning("Selecione variáveis X e Y para calcular correlações.")

# ── ABA 4: EXPORTAR ────────────────────────────────────────────────────────
with tab_export:
    st.markdown('<div class="section-header">Exportar Resultados</div>', unsafe_allow_html=True)
    st.markdown("Todas as figuras também podem ser baixadas individualmente em cada aba (PNG, 300 dpi).")
    st.markdown("---")

    e1, e2 = st.columns(2)

    with e1:
        st.markdown("#### 📗 Excel")
        buf_xl = BytesIO()
        with pd.ExcelWriter(buf_xl, engine="openpyxl") as writer:
            if 'cols_desc' in dir() and cols_desc:
                df[cols_desc].describe().round(4).to_excel(writer, sheet_name="Descritivas")
            if 'df_sh' in dir():
                df_sh.to_excel(writer, sheet_name="Shapiro-Wilk")
            if 'mat_r' in dir():
                mat_r.to_excel(writer, sheet_name="Correlacoes")
                mat_p.to_excel(writer, sheet_name="P-valores")
                mat_n.to_excel(writer, sheet_name="N_observacoes")
                tab_fmt.to_excel(writer, sheet_name="Tabela_Formatada")
        buf_xl.seek(0)
        st.download_button("⬇️ Baixar Excel", buf_xl, file_name="resultados.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    with e2:
        st.markdown("#### 📘 Word")
        try:
            import docx as _docx_check
            docx_ok = True
        except ImportError:
            docx_ok = False

        if docx_ok:
            figs = st.session_state.get("figs", {})
            if figs.get("desc") or figs.get("barras"):
                buf_w = gerar_word(figs, metodo if 'metodo' in dir() else 'pearson',
                                   alpha_c if 'alpha_c' in dir() else 0.05)
                st.download_button("⬇️ Baixar Word", buf_w, file_name="relatorio.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.info("Preencha as abas Descritivas e Gráfico para gerar o relatório.")
        else:
            st.warning("Adicione `python-docx` ao requirements.txt para exportar Word.")

# ── Rodapé ──────────────────────────────────────────────────────────────────
st.markdown("---")
st.caption("Análise Estatística · Descritivas · Gráficos · Correlações · Export Word/Excel/PNG")
